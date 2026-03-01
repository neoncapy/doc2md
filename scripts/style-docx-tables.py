#!/usr/bin/env python3
"""
Post-process pandoc .docx files with professional styling.

Tables:
- Header: Dark Blue (#022366) background, white bold text
- Alternating rows: white / medium blue-grey (#CBD5E1)
- Borders: dark grey (#5A5A5A), 1pt on all lines
- Column widths: content-proportional (sqrt-weighted)

Code blocks (prompts):
- Background: dark slate (#1E293B)
- Text: bright navy (#3B82F6) - unique colour not used elsewhere
"""

import sys
import math
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ── Colour palette: tables ──
HEADER_BG = "022366"
HEADER_TEXT = RGBColor(255, 255, 255)
ALT_ROW = "CBD5E1"
WHITE = "FFFFFF"
BORDER = "5A5A5A"
BORDER_SZ = "8"  # 1pt

# ── Colour palette: code blocks / prompts ──
CODE_BG = "1E293B"           # Dark slate
CODE_TEXT = RGBColor(0x3B, 0x82, 0xF6)  # Bright navy #3B82F6


def style_table_borders(table):
    """Dark grey borders on entire table grid."""
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        table._tbl.insert(0, tblPr)

    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)

    tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{BORDER_SZ}" w:space="0" w:color="{BORDER}"/>'
        f'<w:left w:val="single" w:sz="{BORDER_SZ}" w:space="0" w:color="{BORDER}"/>'
        f'<w:bottom w:val="single" w:sz="{BORDER_SZ}" w:space="0" w:color="{BORDER}"/>'
        f'<w:right w:val="single" w:sz="{BORDER_SZ}" w:space="0" w:color="{BORDER}"/>'
        f'<w:insideH w:val="single" w:sz="{BORDER_SZ}" w:space="0" w:color="{BORDER}"/>'
        f'<w:insideV w:val="single" w:sz="{BORDER_SZ}" w:space="0" w:color="{BORDER}"/>'
        f'</w:tblBorders>'
    ))


def set_cell_shading(cell, color):
    """Set cell background colour."""
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    tcPr.append(parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'
    ))


def set_cell_borders(cell):
    """Borders on individual cell for compatibility."""
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    tcPr.append(parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{BORDER_SZ}" w:space="0" w:color="{BORDER}"/>'
        f'<w:left w:val="single" w:sz="{BORDER_SZ}" w:space="0" w:color="{BORDER}"/>'
        f'<w:bottom w:val="single" w:sz="{BORDER_SZ}" w:space="0" w:color="{BORDER}"/>'
        f'<w:right w:val="single" w:sz="{BORDER_SZ}" w:space="0" w:color="{BORDER}"/>'
        f'</w:tcBorders>'
    ))


def set_fixed_layout(table):
    """Fixed layout so column widths are respected."""
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        table._tbl.insert(0, tblPr)

    for old in tblPr.findall(qn('w:tblLayout')):
        tblPr.remove(old)
    tblPr.append(parse_xml(
        f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>'
    ))


def set_table_width_100pct(table):
    """Set table width to 100% of page."""
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        table._tbl.insert(0, tblPr)

    for old in tblPr.findall(qn('w:tblW')):
        tblPr.remove(old)
    tblPr.append(parse_xml(
        f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>'
    ))


def set_cell_margins(table):
    """Add comfortable padding inside cells."""
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        table._tbl.insert(0, tblPr)

    for old in tblPr.findall(qn('w:tblCellMar')):
        tblPr.remove(old)
    # top/bottom: 40 twips (~0.7mm), left/right: 80 twips (~1.4mm)
    tblPr.append(parse_xml(
        f'<w:tblCellMar {nsdecls("w")}>'
        f'<w:top w:w="40" w:type="dxa"/>'
        f'<w:left w:w="80" w:type="dxa"/>'
        f'<w:bottom w:w="40" w:type="dxa"/>'
        f'<w:right w:w="80" w:type="dxa"/>'
        f'</w:tblCellMar>'
    ))


def calc_widths(table, total_cm=16.0):
    """Content-proportional column widths using sqrt weighting."""
    rows = table.rows
    nr = len(rows)
    nc = len(table.columns)
    if nr < 1 or nc < 1:
        return None

    totals = [0] * nc
    for row in rows:
        cells = row.cells
        for i in range(min(len(cells), nc)):
            totals[i] += len(cells[i].text.strip())

    avgs = [t / nr for t in totals]

    # Header lengths
    hdr = rows[0].cells
    hdr_lens = [
        len(hdr[i].text.strip()) if i < len(hdr) else 0
        for i in range(nc)
    ]

    # sqrt-weighted: prevents extreme skew
    weights = [
        math.sqrt(max(avgs[i], hdr_lens[i], 3))
        for i in range(nc)
    ]

    tw = sum(weights)
    widths = [(w / tw) * total_cm for w in weights]

    # Minimum 1.5cm per column
    for i in range(nc):
        widths[i] = max(widths[i], 1.5)

    # Re-normalise
    s = sum(widths)
    return [(w / s) * total_cm for w in widths]


def set_paragraph_shading(paragraph, color):
    """Set background colour on a paragraph."""
    pPr = paragraph._element.get_or_add_pPr()
    for old in pPr.findall(qn('w:shd')):
        pPr.remove(old)
    pPr.append(parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'
    ))


def style_code_blocks(doc):
    """Style Source Code paragraphs (prompts, code blocks).
    Dark slate background, bright navy text."""
    count = 0
    for para in doc.paragraphs:
        if para.style.name == 'Source Code':
            count += 1
            # Background
            set_paragraph_shading(para, CODE_BG)
            # Text colour on every run
            for run in para.runs:
                run.font.color.rgb = CODE_TEXT
    return count


def style_docx(input_path, output_path):
    """Apply professional styling to tables and code blocks."""
    doc = Document(input_path)
    table_count = 0

    # ── Tables ──
    for table in doc.tables:
        nr = len(table.rows)
        nc = len(table.columns)
        if nr == 0 or nc == 0:
            continue
        table_count += 1

        # Table-level properties
        style_table_borders(table)
        set_fixed_layout(table)
        set_table_width_100pct(table)
        set_cell_margins(table)

        # Column widths
        widths = calc_widths(table)
        if widths:
            for i in range(min(nc, len(widths))):
                table.columns[i].width = Cm(widths[i])
            for row in table.rows:
                cells = row.cells
                for i in range(min(len(cells), nc, len(widths))):
                    cells[i].width = Cm(widths[i])

        # Cell-level styling
        for row_idx, row in enumerate(table.rows):
            cells = row.cells
            for col_idx in range(min(len(cells), nc)):
                cell = cells[col_idx]

                # Borders per cell (redundant safety)
                set_cell_borders(cell)

                # Background colour
                if row_idx == 0:
                    set_cell_shading(cell, HEADER_BG)
                elif row_idx % 2 == 0:
                    set_cell_shading(cell, ALT_ROW)
                else:
                    set_cell_shading(cell, WHITE)

                # Text formatting (header only)
                for para in cell.paragraphs:
                    for run in para.runs:
                        if row_idx == 0:
                            run.font.color.rgb = HEADER_TEXT
                            run.font.bold = True

    # ── Code blocks / prompts ──
    code_count = style_code_blocks(doc)

    doc.save(output_path)
    print(f"Styled {table_count} tables, {code_count} code blocks -> {output_path}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print(f"Usage: python {sys.argv[0]} input.docx output.docx")
        sys.exit(1)
    style_docx(sys.argv[1], sys.argv[2])
